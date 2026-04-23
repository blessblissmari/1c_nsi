import shutil
import subprocess
import tempfile
from pathlib import Path

import openpyxl
from loguru import logger


def _clean(val):
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return val
    s = str(val).replace("\u00a0", " ").strip()
    if not s:
        return None
    return s


def parse_xlsx(file_path: str | Path, sheet_name: str | None = None, header_row: int = 0) -> list[dict]:
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [_clean(h) or f"col_{i}" for i, h in enumerate(rows[header_row])]
    data = []

    for row in rows[header_row + 1 :]:
        item = {}
        for i, val in enumerate(row):
            if i < len(headers):
                item[headers[i]] = _clean(val) if val is not None else val
        if any(v is not None for v in row):
            data.append(item)

    wb.close()
    logger.info(f"Parsed {len(data)} rows from {file_path}")
    return data


def parse_pdf_text(file_path: str | Path) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    result = "\n".join(text_parts)
    logger.info(f"Parsed PDF text from {file_path}: {len(result)} chars")
    return result


def parse_pdf_scan(file_path: str | Path) -> str:
    # Prefer local MinerU pipeline extraction/OCR (no external API key required).
    # We use MinerU for both "text PDFs" and scanned PDFs to keep behavior consistent.
    mineru_bin = shutil.which("mineru")
    if mineru_bin:
        try:
            with tempfile.TemporaryDirectory(prefix="mineru_out_") as outdir:
                cmd = [
                    mineru_bin,
                    "-p",
                    str(file_path),
                    "-o",
                    outdir,
                    "-b",
                    "pipeline",
                    "-m",
                    "ocr",
                    "-l",
                    "cyrillic",
                    "--formula",
                    "false",
                    "--table",
                    "false",
                ]
                logger.info(f"Running MinerU pipeline OCR: {' '.join(cmd[:6])} ...")
                subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=900)

                md_files = sorted(Path(outdir).rglob("*.md"), key=lambda p: p.stat().st_size, reverse=True)
                if md_files:
                    text = md_files[0].read_text(encoding="utf-8", errors="ignore")
                    if text.strip():
                        logger.info(f"MinerU produced markdown: {md_files[0].name} ({len(text)} chars)")
                        return text

                txt_files = sorted(Path(outdir).rglob("*.txt"), key=lambda p: p.stat().st_size, reverse=True)
                if txt_files:
                    text = txt_files[0].read_text(encoding="utf-8", errors="ignore")
                    if text.strip():
                        logger.info(f"MinerU produced text: {txt_files[0].name} ({len(text)} chars)")
                        return text
        except subprocess.TimeoutExpired:
            logger.warning("MinerU OCR timed out; falling back to pdfplumber")
        except Exception as e:
            logger.warning(f"MinerU OCR failed: {e}; falling back to pdfplumber")
    else:
        logger.warning("MinerU CLI not found on PATH; skipping OCR and falling back to pdfplumber")

    return parse_pdf_text(file_path)


def parse_docx(file_path: str | Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            tables_data.append(" | ".join(row_data))

    result = "\n".join(paragraphs)
    if tables_data:
        result += "\n\n--- Таблицы ---\n" + "\n".join(tables_data)

    logger.info(f"Parsed DOCX from {file_path}: {len(result)} chars")
    return result


def parse_file(file_path: str | Path) -> str | list[dict]:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".xlsx":
        return parse_xlsx(file_path)
    elif ext == ".pdf":
        # Always route PDFs through MinerU first, fallback to pdfplumber.
        return parse_pdf_scan(file_path)
    elif ext in {".png", ".jpg", ".jpeg"}:
        # MinerU can process images via pipeline as well.
        return parse_pdf_scan(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".txt":
        return path.read_text(encoding="utf-8")
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return ""


def _try_pdf(file_path: str | Path) -> str:
    text_result = parse_pdf_text(file_path)
    if text_result.strip():
        return text_result
    logger.info("PDF appears to be scanned, trying OCR")
    return parse_pdf_scan(file_path)


def detect_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    type_map = {
        ".xlsx": "excel",
        ".xls": "excel",
        ".pdf": "pdf",
        ".docx": "word",
        ".doc": "word",
        ".txt": "text",
        ".csv": "csv",
    }
    return type_map.get(ext, "unknown")
